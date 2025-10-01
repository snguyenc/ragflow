#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
import re
from functools import reduce
from io import BytesIO
from timeit import default_timer as timer

from docx import Document
from docx.image.exceptions import InvalidImageStreamError, UnexpectedEndOfFileError, UnrecognizedImageError
from docx.opc.pkgreader import _SerializedRelationships, _SerializedRelationship
from docx.opc.oxml import parse_xml
from markdown import markdown
from PIL import Image

from deepdoc.parser import DocxParser, MarkdownElementExtractor, MarkdownParser, PdfParser, TxtParser
from rag.nlp import concat_img, find_codec, naive_merge, naive_merge_with_images, naive_merge_docx, rag_tokenizer, tokenize_chunks, tokenize_chunks_with_images, tokenize_table
from rag.svr.bookstack_svr import fetch_bookstack_pages

class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        img = paragraph._element.xpath('.//pic:pic')
        if not img:
            return None
        img = img[0]
        embed = img.xpath('.//a:blip/@r:embed')
        if not embed:
            return None
        embed = embed[0]
        try:
            related_part = document.part.related_parts[embed]
            image_blob = related_part.image.blob
        except UnrecognizedImageError:
            logging.info("Unrecognized image format. Skipping image.")
            return None
        except UnexpectedEndOfFileError:
            logging.info("EOF was unexpectedly encountered while reading an image stream. Skipping image.")
            return None
        except InvalidImageStreamError:
            logging.info("The recognized image stream appears to be corrupted. Skipping image.")
            return None
        except UnicodeDecodeError:
            logging.info("The recognized image stream appears to be corrupted. Skipping image.")
            return None
        except Exception:
            logging.info("The recognized image stream appears to be corrupted. Skipping image.")
            return None
        try:
            image = Image.open(BytesIO(image_blob)).convert('RGB')
            return image
        except Exception:
            return None

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __get_nearest_title(self, table_index, filename):
        """Get the hierarchical title structure before the table"""
        import re
        from docx.text.paragraph import Paragraph

        titles = []
        blocks = []

        # Get document name from filename parameter
        doc_name = re.sub(r"\.[a-zA-Z]+$", "", filename)
        if not doc_name:
            doc_name = "Untitled Document"

        # Collect all document blocks while maintaining document order
        try:
            # Iterate through all paragraphs and tables in document order
            for i, block in enumerate(self.doc._element.body):
                if block.tag.endswith('p'):  # Paragraph
                    p = Paragraph(block, self.doc)
                    blocks.append(('p', i, p))
                elif block.tag.endswith('tbl'):  # Table
                    blocks.append(('t', i, None))  # Table object will be retrieved later
        except Exception as e:
            logging.error(f"Error collecting blocks: {e}")
            return ""

        # Find the target table position
        target_table_pos = -1
        table_count = 0
        for i, (block_type, pos, _) in enumerate(blocks):
            if block_type == 't':
                if table_count == table_index:
                    target_table_pos = pos
                    break
                table_count += 1

        if target_table_pos == -1:
            return ""  # Target table not found

        # Find the nearest heading paragraph in reverse order
        nearest_title = None
        for i in range(len(blocks)-1, -1, -1):
            block_type, pos, block = blocks[i]
            if pos >= target_table_pos:  # Skip blocks after the table
                continue

            if block_type != 'p':
                continue

            if block.style and block.style.name and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                try:
                    level_match = re.search(r"(\d+)", block.style.name)
                    if level_match:
                        level = int(level_match.group(1))
                        if level <= 7:  # Support up to 7 heading levels
                            title_text = block.text.strip()
                            if title_text:  # Avoid empty titles
                                nearest_title = (level, title_text)
                                break
                except Exception as e:
                    logging.error(f"Error parsing heading level: {e}")

        if nearest_title:
            # Add current title
            titles.append(nearest_title)
            current_level = nearest_title[0]

            # Find all parent headings, allowing cross-level search
            while current_level > 1:
                found = False
                for i in range(len(blocks)-1, -1, -1):
                    block_type, pos, block = blocks[i]
                    if pos >= target_table_pos:  # Skip blocks after the table
                        continue

                    if block_type != 'p':
                        continue

                    if block.style and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                        try:
                            level_match = re.search(r"(\d+)", block.style.name)
                            if level_match:
                                level = int(level_match.group(1))
                                # Find any heading with a higher level
                                if level < current_level:
                                    title_text = block.text.strip()
                                    if title_text:  # Avoid empty titles
                                        titles.append((level, title_text))
                                        current_level = level
                                        found = True
                                        break
                        except Exception as e:
                            logging.error(f"Error parsing parent heading: {e}")

                if not found:  # Break if no parent heading is found
                    break

            # Sort by level (ascending, from highest to lowest)
            titles.sort(key=lambda x: x[0])
            # Organize titles (from highest to lowest)
            hierarchy = [doc_name] + [t[1] for t in titles]
            return " > ".join(hierarchy)

        return ""

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        last_image = None
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page:
                if p.text.strip():
                    if p.style and p.style.name == 'Caption':
                        former_image = None
                        if lines and lines[-1][1] and lines[-1][2] != 'Caption':
                            former_image = lines[-1][1].pop()
                        elif last_image:
                            former_image = last_image
                            last_image = None
                        lines.append((self.__clean(p.text), [former_image], p.style.name))
                    else:
                        current_image = self.get_picture(self.doc, p)
                        image_list = [current_image]
                        if last_image:
                            image_list.insert(0, last_image)
                            last_image = None
                        lines.append((self.__clean(p.text), image_list, p.style.name if p.style else ""))
                else:
                    if current_image := self.get_picture(self.doc, p):
                        if lines:
                            lines[-1][1].append(current_image)
                        else:
                            last_image = current_image
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        new_line = [(line[0], reduce(concat_img, line[1]) if line[1] else None) for line in lines]

        tbls = []
        for i, tb in enumerate(self.doc.tables):
            title = self.__get_nearest_title(i, filename)
            html = "<table>"
            if title:
                html += f"<caption>Table Location: {title}</caption>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                try:
                    while i < len(r.cells):
                        span = 1
                        c = r.cells[i]
                        for j in range(i + 1, len(r.cells)):
                            if c.text == r.cells[j].text:
                                span += 1
                                i = j
                            else:
                                break
                        i += 1
                        html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                except Exception as e:
                    logging.warning(f"Error parsing table, ignore: {e}")
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return new_line, tbls


class Pdf(PdfParser):
    def __init__(self):
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None, separate_tables_figures=False):
        start = timer()
        first_start = start
        callback(msg="OCR started")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.info("OCR({}~{}): {:.2f}s".format(from_page, to_page, timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.63, "Layout analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge()
        callback(0.67, "Text merged ({:.2f}s)".format(timer() - start))

        if separate_tables_figures:
            tbls, figures = self._extract_table_figure(True, zoomin, True, True, True)
            self._concat_downward()
            logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls, figures
        else:
            tbls = self._extract_table_figure(True, zoomin, True, True)
            self._naive_vertical_merge()
            self._concat_downward()
            # self._filter_forpages()
            logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls


class Markdown(MarkdownParser):
    def get_picture_urls(self, sections):
        if not sections:
            return []
        if isinstance(sections, type("")):
            text = sections
        elif isinstance(sections[0], type("")):
            text = sections[0]
        else:
            return []

        from bs4 import BeautifulSoup
        html_content = markdown(text)
        soup = BeautifulSoup(html_content, 'html.parser')
        html_images = [img.get('src') for img in soup.find_all('img') if img.get('src')]
        return html_images

    def get_pictures(self, text):
        """Download and open all images from markdown text."""
        import requests
        image_urls = self.get_picture_urls(text)
        images = []
        # Find all image URLs in text
        for url in image_urls:
            try:
                # check if the url is a local file or a remote URL
                if url.startswith(('http://', 'https://')):
                    # For remote URLs, download the image
                    response = requests.get(url, stream=True, timeout=30)
                    if response.status_code == 200 and response.headers['Content-Type'].startswith('image/'):
                        img = Image.open(BytesIO(response.content)).convert('RGB')
                        images.append(img)
                else:
                    # For local file paths, open the image directly
                    from pathlib import Path
                    local_path = Path(url)
                    if not local_path.exists():
                        logging.warning(f"Local image file not found: {url}")
                        continue
                    img = Image.open(url).convert('RGB')
                    images.append(img)
            except Exception as e:
                logging.error(f"Failed to download/open image from {url}: {e}")
                continue

        return images if images else None

    def __call__(self, filename, binary=None, separate_tables=True):
        if binary:
            encoding = find_codec(binary)
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r") as f:
                txt = f.read()

        remainder, tables = self.extract_tables_and_remainder(f'{txt}\n', separate_tables=separate_tables)

        extractor = MarkdownElementExtractor(txt)
        element_sections = extractor.extract_elements()
        sections = [(element, "") for element in element_sections]

        tbls = []
        for table in tables:
            tbls.append(((None, markdown(table, extensions=['markdown.extensions.tables'])), ""))
        return sections, tbls

def load_from_xml_v2(baseURI, rels_item_xml):
    """
    Return |_SerializedRelationships| instance loaded with the
    relationships contained in *rels_item_xml*. Returns an empty
    collection if *rels_item_xml* is |None|.
    """
    srels = _SerializedRelationships()
    if rels_item_xml is not None:
        rels_elm = parse_xml(rels_item_xml)
        for rel_elm in rels_elm.Relationship_lst:
            if rel_elm.target_ref in ('../NULL', 'NULL'):
                continue
            srels._srels.append(_SerializedRelationship(baseURI, rel_elm))
    return srels

def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
        Supported file formats are docx, pdf, excel, txt, and BookStack documents.
        For BookStack documents, content is fetched from BookStack API instead of local files.
        This method apply the naive ways to chunk files.
        Successive text will be sliced into pieces using 'delimiter'.
        Next, these successive pieces are merge into chunks whose token number is no more than 'Max token number'.
    """

    is_english = lang.lower() == "english"  # is_english(cks)
    parser_config = kwargs.get(
        "parser_config", {
            "chunk_token_num": 512, "delimiter": "\n!?。；！？", "layout_recognize": "DeepDOC"})
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    pdf_parser = None
    section_images = None
    sections = []
    # Handle BookStack documents
    callback(0.1, "Start to fetch BookStack content.")
    try:
        pages = fetch_bookstack_pages(parser_config, kwargs, callback)
        callback(0.8, "Finish fetching BookStack content.")

        for page in pages:
            #logging.info(f"BookStack page content: {page.content}")
            metadata = page.metadata
            doc["category_kwd"] = metadata.get("category", "")
            doc["guide_kwd"] = metadata.get("guide", "")
            doc["page_id"] = metadata.get("page_id", "")
            doc["revision_count"] = metadata.get("revision_count", 1)
            doc["book_id"] = metadata.get("book_id", "")
            doc["chapter_id"] = metadata.get("chapter_id", "")
            
            tags_map = {tag["name"]: tag["value"] for tag in metadata.get("tags", [])}
            doc["article_type_kwd"] = tags_map.get("ArticleType", "Topic")

            sections.append(page.content)

        callback(0.9, "Finish fetching BookStack content.")
    except Exception as e:
        callback(-1, f"BookStack fetch error: {str(e)}")
        logging.error(f"BookStack fetch error: {str(e)}")
        return []
    

    st = timer()
    if section_images:
        # if all images are None, set section_images to None
        if all(image is None for image in section_images):
            section_images = None

    if section_images:
        chunks, images = naive_merge_with_images(sections, section_images,
                                        int(parser_config.get(
                                            "chunk_token_num", 128)), parser_config.get(
                                            "delimiter", "\n!?。；！？"))
        if kwargs.get("section_only", False):
            return chunks

        res.extend(tokenize_chunks_with_images(chunks, doc, is_english, images))
    else:
        chunks = naive_merge(
            sections, int(parser_config.get(
                "chunk_token_num", 128)), parser_config.get(
                "delimiter", "\n!?。；！？"))
        res.extend(tokenize_chunks(chunks, doc, is_english, pdf_parser))

    logging.info("naive_merge({}): {}".format(filename, timer() - st))
    return res


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass

    chunk("Chapter: abcd", from_page=0, to_page=10, callback=dummy)
